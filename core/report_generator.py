from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document
import os

def generate_report(capture_file, format):
    if format == 'pdf':
        return generate_pdf_report(capture_file)
    elif format == 'txt':
        return generate_txt_report(capture_file)
    elif format == 'xlsx':
        return generate_excel_report(capture_file)
    elif format == 'docx':
        return generate_word_report(capture_file)
    else:
        raise ValueError("Unsupported format")

def generate_pdf_report(capture_file):
    c = canvas.Canvas("report.pdf", pagesize=letter)
    c.drawString(100, 750, "PacketCatcher Report")
    c.drawString(100, 730, f"Capture File: {capture_file}")
    c.save()
    return "report.pdf"

def generate_txt_report(capture_file):
    with open("report.txt", "w") as f:
        f.write(f"PacketCatcher Report\n")
        f.write(f"Capture File: {capture_file}\n")
    return "report.txt"

def generate_excel_report(capture_file):
    wb = Workbook()
    ws = wb.active
    ws.title = "PacketCatcher Report"
    ws['A1'] = "PacketCatcher Report"
    ws['A2'] = f"Capture File: {capture_file}"
    wb.save("report.xlsx")
    return "report.xlsx"

def generate_word_report(capture_file):
    doc = Document()
    doc.add_heading('PacketCatcher Report', level=1)
    doc.add_paragraph(f'Capture File: {capture_file}')
    doc.save('report.docx')
    return "report.docx"